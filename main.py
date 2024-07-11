import os, subprocess
import datetime as dt
import tkinter as tk
import traceback
from tkinter import filedialog
from tkinter import messagebox
import docx
from docx2pdf import convert
from tkinter import Button as tkButton


class InvoiceAutomation:
    def __init__(self):
       # self.create_invoice = None
        self.root = tk.Tk() # Create a new Tkinter window
        self.root.title('Invoice Automation')
        self.root.geometry('700x900')



        # Create the labels and entries for the invoice automation form
        self.partner_label = tk.Label(self.root, text='Location Name:')
        self.partner_entry = tk.Entry(self.root)

        self.partner_street_label = tk.Label(self.root, text='Partner Street:')
        self.partner_street_entry = tk.Entry(self.root)

        self.partner_zip_city_country_label = tk.Label(self.root, text='Partner Zip, City, Country:')
        self.partner_zip_city_country_entry = tk.Entry(self.root)

        self.invoice_number_label = tk.Label(self.root, text='Invoice Number:')
        self.invoice_number_entry = tk.Entry(self.root)

        self.invoice_date_label = tk.Label(self.root, text='Invoice Date:')
        self.invoice_date_entry = tk.Entry(self.root)

        self.service_description_label = tk.Label(self.root, text='Service Description:')
        self.service_description_entry = tk.Entry(self.root)

        self.service_amount_label = tk.Label(self.root, text='Service Amount:')
        self.service_amount_entry = tk.Entry(self.root)

        self.location_name_label = tk.Label(self.root, text='Location Name:')
        self.location_name_entry = tk.Entry(self.root)

        self.service_single_amount_label = tk.Label(self.root, text='Service Single Amount:')
        self.service_single_amount_entry = tk.Entry(self.root)

        self.payment_terms_label = tk.Label(self.root, text='Payment Terms:')
        self.payment_terms_entry = tk.Entry(self.root)

        self.payment_method_label = tk.Label(self.root, text='Payment Method:')

        self.schduled_date_label = tk.Label(self.root, text='Scheduled Date:')
        self.scheduled_date = tk.Entry(self.root)

        self.provider_name_label = tk.Label(self.root, text='Provider Name:')
        self.provider_name = tk.Entry(self.root)

        self.created_by_label = tk.Label(self.root, text='Created By:')
        self.created_by = tk.Entry(self.root)

        self.notes_label = tk.Label(self.root, text='Notes:')
        self.notes = tk.Entry(self.root)

        self.notes_createdby_label = tk.Label(self.root, text='Notes Created By:')
        self.notes_createdby = tk.Entry(self.root)

        self.notes_data_label = tk.Label(self.root, text='Notes Data:')
        self.notes_data = tk.Entry(self.root)

       # Create the button to create the invoice
        self.nte = tk.Label(self.root, text='NTE:')
        self.nte_entry = tk.Entry(self.root)




    # Create the dropdown for the payment method
        self.payment_method_entry = tk.Entry(self.root)
        self.payment_methods = {'Credit Card', 'Check', 'Wire Transfer'}
        self.payment_methods = tk.StringVar(self.root)
        self.payment_methods.set('Wire Transfer')
        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_methods, "Credit Card", "Check", "Wire Transfer")

        self.create_button_label = tk.Label(self.root, text='Create Invoice')
        self.create_button = tk.Entry(self.root)

        padding_options = {'fill': tk.X, 'expand': True, 'padx': 5, 'pady': 5}

        self.tkbutton = tkButton(self.root, text='Create Invoice', command=self.create_invoice)
        self.tkbutton.pack(padding_options)

    # Pack the widgets  in the window


        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)

        self.partner_street_label.pack(padding_options)
        self.partner_street_entry.pack(padding_options)

        self.partner_zip_city_country_label.pack(padding_options)
        self.partner_zip_city_country_entry.pack(padding_options)

        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)

        self.invoice_date_label.pack(padding_options)
        self.invoice_date_entry.pack(padding_options)

        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)

        self.service_amount_label.pack(padding_options)
        self.service_amount_entry.pack(padding_options)

        self.service_single_amount_label.pack(padding_options)
        self.service_single_amount_entry.pack(padding_options)

        self.payment_terms_label.pack(padding_options)
        self.payment_terms_entry.pack(padding_options)

        self.payment_method_label.pack(padding_options)

        self.payment_method_entry.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)

        self.nte.pack(padding_options)
        self.nte_entry.pack(padding_options)

        self.location_name_label.pack(padding_options)
        self.location_name_entry.pack(padding_options)

        self.schduled_date_label.pack(padding_options)
        self.scheduled_date.pack(padding_options)

        self.provider_name_label.pack(padding_options)
        self.provider_name.pack(padding_options)

        self.created_by_label.pack(padding_options)
        self.created_by.pack(padding_options)

        self.notes_label.pack(padding_options)
        self.notes.pack(padding_options)

        self.notes_createdby_label.pack(padding_options)
        self.notes_createdby.pack(padding_options)

        self.notes_data_label.pack(padding_options)
        self.notes_data.pack(padding_options)




    #self.create_button.pack(padding_options)
        #self.create_button_label.pack(padding_options)

        # self.create_button.pack(padding_options)

        self.root.mainloop() # Start the main loop



    # Create the invoice, save it to the file system, and open it
    @staticmethod
    def replace_text(doc, old_text, new_text):
        for p in doc.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)

    @staticmethod
    def save_invoice(doc):
        # logic to save files to the file system
        #save_file = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF files', '*.pdf')]) # Ask the user to save the file
        save_file = filedialog.asksaveasfilename(defaultextension='.docx', filetypes=[('DOCX files', '*.docx')])

        doc.save(save_file) # Save the docx file
        convert(save_file) # Convert the docx file to a pdf file
        convert(save_file.replace('.docx', '.pdf')) # Convert the docx file to a pdf file
        os.rename(save_file.replace('.docx', '.pdf'), save_file) # Rename the file to a pdf file
        subprocess.Popen(['open', save_file])
        print(f'[+] Invoice saved to {save_file} \n {os.getcwd()}')
        return f'[+] {save_file}  Invoice saved to {save_file} \n {os.getcwd()}'

    # loads docx template file,
    def create_invoice(self):
        doc = docx.Document('template.docx')
        selected_payment_method = self.payment_methods.get()
        print("[!] payment mehtods selected: ", selected_payment_method)

        # replace data using a dictionary for
        try:
            replacements = {
                "<<Location.Name>>": self.partner_entry.get(),
                "<<Street>>": self.partner_street_entry.get(),
                "<<ZipCityCountry>>": self.partner_zip_city_country_entry.get(),
                "<<InvoiceNumber>>": self.invoice_number_entry.get(),
                "<<InvoiceDate>>": self.invoice_date_entry.get(),
                "<<ServiceAmount>>": self.service_amount_entry.get(),
                "<<ServiceSingleAmount>>": self.service_single_amount_entry.get(),
                "<<PaymentTerms>>": self.payment_terms_entry.get(),
                "<<PaymentMethod>>": selected_payment_method,
                "<<ScheduledDate>>": self.scheduled_date.get(),
                "<<ProviderName>>": self.provider_name.get(),
                "<<NTE>>": self.nte_entry.get(),
                "<<LocationName>>": self.location_name_entry.get(),
                "<<Created By>>": self.created_by.get(),
                "<<Notes>>": self.notes.get(),
                "<<Notes Created By>>": self.notes_createdby.get(),
                "<<Notes Data>>": self.notes_data.get()

            }

        except ValueError as e:
            messagebox.showerror('Error', f'Error replacing placeholders: {e} \n {traceback.format_exc()}')
            print(f'Error replacing placeholders:  {e}')
            print(traceback.format_exc())
            return

        # Save the invoice to the file system
        for paragraph in list(doc.paragraphs):
            print(paragraph.text)
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    if paragraph.text:
                        print("REPLACE TEXT TEST CASE: \n", paragraph.text)
                    self.replace_text(doc, key, value) # calls static method to replace text in the docx file

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                if paragraph.text:
                                    print("REPLACE TEXT TEST CASE: \n", paragraph.text)
                                self.replace_text(doc, key, value)

        save_invoice = self.save_invoice(doc)
        print(save_invoice)

    # Replace the placeholders with data from internal program
    ## -->> NEED TO SET A GETTER/SETTER FOR THE DATA SO IT CAN BE PARSED FROM THE MAIN APP <<--
    def create_invoice_INTERNALDATA(self):
        doc = docx.Document('invoice_template.docx')
        selected_payment_method = self.payment_methods.get()
        print("[!] payment mehtods selected: ", selected_payment_method)

        try:
            self.replace_text(doc, '<<Partner>>', self.partner_entry.get())
            self.replace_text(doc, '<<PartnerStreet>>', self.partner_street_entry.get())
            self.replace_text(doc, '<<PartnerZipCityCountry>>', self.partner_zip_city_country_entry.get())
            self.replace_text(doc, '<<InvoiceNumber>>', self.invoice_number_entry.get())
            self.replace_text(doc, '<<InvoiceDate>>', self.invoice_date_entry.get())
            self.replace_text(doc, '<<ServiceDescription>>', self.service_description_entry.get())
            self.replace_text(doc, '<<ServiceAmount>>', self.service_amount_entry.get())
            self.replace_text(doc, '<<ServiceSingleAmount>>', self.service_single_amount_entry.get())
            self.replace_text(doc, '<<PaymentTerms>>', self.payment_terms_entry.get())
            self.replace_text(doc, '<<PaymentMethod>>', selected_payment_method)
            self.replace_text(doc, '<<ScheduledDate>>', self.scheduled_date.get())
            self.replace_text(doc, '<<ProviderName>>', self.provider_name.get())
            self.replace_text(doc, '<<NTE>>', self.nte_entry.get())
            self.replace_text(doc, '<<LocationName>>', self.location_name_entry.get())
            self.replace_text(doc, '<<PaymentMethod>>', selected_payment_method)
            self.replace_text(doc, '<<ScheduledDate>>', self.scheduled_date.get())
            self.replace_text(doc, '<<Created By>>', self.created_by.get())
            self.replace_text(doc, '<<Notes>>', self.notes.get())
            self.replace_text(doc, '<<Notes Created By>>', self.notes_createdby.get())
            self.replace_text(doc, '<<Notes Data>>', self.notes_data.get())
            self.replace_text(doc, '<<NTE>>', self.nte_entry.get())

        except ValueError as e:
            messagebox.showerror('Error', f'Error replacing placeholders: {e} \n {traceback.format_exc()}')
            print(f'Error replacing placeholders:  {e}')
            print(traceback.format_exc())
            return




if __name__ == '__main__':
    InvoiceAutomation()
    # invoice_automation = InvoiceAutomation()
    # invoice_automation.root.mainloop()
    #

